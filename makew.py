from docx import Document
from docx.shared import Inches

guestnames=['Ivan', 'Oleg', 'Olga', 'Maria and Alexander']

for x in guestnames:
    document = Document()
    document.add_heading('To Wedding', 0)
    p = document.add_paragraph(x+', we glad to invite you to wedding') 
    p.add_run(' of Tatyana and Dmitry Ratuenko').bold = True
    p.add_run(', in 12.05.2017 in Silent Rosha Usatba.')
    document.add_heading('We waiting for you!!!', level=1)
    document.add_paragraph('')
    document.add_page_break()
    document.save(x+'.docx')